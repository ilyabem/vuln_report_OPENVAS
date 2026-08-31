#!/usr/bin/env python3
"""
vuln_report.py

Объединенный инструмент: обрабатывает XML-отчеты Greenbone OpenVAS,
группирует уязвимости по OID и выдает результат в любом сочетании
форматов: JSON (структурированные данные), DOCX (человекочитаемый
Word-документ) и PDF (через LibreOffice).

Использование:
    # Только Word-документ (формат по умолчанию, если ничего не указано)
    python vuln_report.py -i /path/to/reports/

    # Word + PDF
    python vuln_report.py -i /path/to/reports/ --docx --pdf

    # Все три формата сразу
    python vuln_report.py -i /path/to/reports/ --json --docx --pdf

    # Только PDF (промежуточный .docx создается и удаляется автоматически)
    python vuln_report.py -i /path/to/reports/ --pdf

    # Рекурсивный обход каталога и отсев информационных записей
    python vuln_report.py -i /path/to/reports/ -r --min-cvss 4.0 --docx

    # Свое базовое имя выходных файлов (без расширения)
    python vuln_report.py -i report.xml -o audit_2026_06 --docx --pdf

    # С водяным знаком (логотип организации на каждой странице)
    python vuln_report.py -i /path/to/reports/ --pdf --watermark /path/to/logo.png

    # Водяной знак с пользовательской шириной и непрозрачностью
    python vuln_report.py -i /path/to/reports/ --docx \\
        --watermark logo.png --watermark-width 6 --watermark-opacity 20

Коды возврата:
    0 — все запрошенные форматы успешно сохранены
    1 — нечего обрабатывать либо часть форматов не создана
    2 — некорректные аргументы командной строки
    130 — прервано пользователем (Ctrl+C)

Зависимости:
    pip install python-docx
    pip install defusedxml   # опционально: защита от XML-бомб, используется если есть
    (для --pdf необходим установленный LibreOffice: soffice в PATH)
    Примечание: непрозрачность водяного знака (--watermark-opacity) работает
    корректно в Microsoft Word; LibreOffice при экспорте в PDF её игнорирует.
"""

import argparse
import gzip
import ipaddress
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import xml.etree.ElementTree as ET
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# defusedxml защищает от entity-бомб ("billion laughs") в недоверенных отчетах.
# Если пакет не установлен — работаем на стандартном ElementTree.
try:
    from defusedxml.ElementTree import iterparse as _safe_iterparse
    HAVE_DEFUSEDXML = True
except ImportError:  # pragma: no cover - зависит от окружения
    _safe_iterparse = ET.iterparse
    HAVE_DEFUSEDXML = False


# ============================================================================
# Общие константы

# XML-пространства имён, используемые в водяном знаке
_WP  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
_A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
_PIC = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
_W   = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
# ============================================================================

THREAT_ORDER = ["Critical", "High", "Medium", "Low", "Log", "None"]

# Ранг для сортировки: чем меньше, тем опаснее
THREAT_RANK = {level: i for i, level in enumerate(THREAT_ORDER)}

THREAT_COLORS = {
    "Critical": "C00000",
    "High": "E53935",
    "Medium": "F57C00",
    "Low": "FBC02D",
    "Log": "9E9E9E",
    "None": "9E9E9E",
}

THREAT_TEXT_LIGHT = {"Critical", "High", "Medium"}  # светлый текст на темном фоне

RU_THREAT = {
    "Critical": "Критический",
    "High": "Высокий",
    "Medium": "Средний",
    "Low": "Низкий",
    "Log": "Инфо",
    "None": "Не определен",
}

# Синонимы уровней угрозы, встречающиеся в разных версиях/локализациях OpenVAS
THREAT_ALIASES = {
    "critical": "Critical",
    "high": "High",
    "medium": "Medium",
    "low": "Low",
    "log": "Log",
    "info": "Log",
    "informational": "Log",
    "debug": "Log",
    "alarm": "High",
    "false positive": "None",
    "none": "None",
    "": "None",
}

XML_SUFFIXES = (".xml", ".xml.gz")

# Символы, недопустимые в XML 1.0: попав в run, они делают .docx нечитаемым
_ILLEGAL_XML_CHARS = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f﷐-﷟￾￿]"
)

QUIET = False


def log(msg):
    """Обычный прогресс-вывод (подавляется флагом --quiet)."""
    if not QUIET:
        print(msg, flush=True)


def warn(msg):
    """Предупреждения и ошибки — всегда в stderr, чтобы не смешивать с данными."""
    print(msg, file=sys.stderr, flush=True)


def normalize_threat(value):
    """Приводит уровень угрозы к одному из ключей THREAT_ORDER."""
    if not value:
        return "None"
    key = str(value).strip().lower()
    if key in THREAT_ALIASES:
        return THREAT_ALIASES[key]
    # "High (CVSS: 7.5)" и подобные варианты
    for alias, canonical in THREAT_ALIASES.items():
        if alias and key.startswith(alias):
            return canonical
    return "None"


def clean_text(value):
    """Убирает управляющие символы, ломающие OOXML, и нормализует переводы строк."""
    if not value:
        return ""
    text = _ILLEGAL_XML_CHARS.sub("", str(value))
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


# ============================================================================
# Часть 1: парсинг XML-отчетов OpenVAS и группировка по OID
# ============================================================================

def _is_xml_name(fname):
    return fname.lower().endswith(XML_SUFFIXES)


def find_xml_files(input_path, recursive=False):
    """Возвращает список путей к XML-файлам (.xml/.xml.gz) для обработки."""
    if not os.path.exists(input_path):
        warn(f"[ОШИБКА] Путь не существует: {input_path}")
        return []

    if os.path.isfile(input_path):
        if _is_xml_name(input_path):
            return [input_path]
        warn(f"[ОШИБКА] Файл не является XML: {input_path}")
        return []

    if os.path.isdir(input_path):
        files = []
        if recursive:
            for dirpath, dirnames, filenames in os.walk(input_path):
                dirnames.sort()
                for fname in sorted(filenames):
                    if _is_xml_name(fname):
                        files.append(os.path.join(dirpath, fname))
        else:
            for fname in sorted(os.listdir(input_path)):
                full = os.path.join(input_path, fname)
                if os.path.isfile(full) and _is_xml_name(fname):
                    files.append(full)
        if not files:
            hint = "" if recursive else " (попробуйте -r для обхода подкаталогов)"
            warn(f"[ПРЕДУПРЕЖДЕНИЕ] В каталоге не найдено XML-файлов: "
                 f"{input_path}{hint}")
        return files

    warn(f"[ОШИБКА] Неподдерживаемый тип пути: {input_path}")
    return []


def text_or_none(elem):
    if elem is None:
        return None
    return clean_text(elem.text) or None


def severity_to_threat(score):
    """Резервное определение уровня угрозы по CVSS, если тег <threat> отсутствует."""
    try:
        score = float(score)
    except (TypeError, ValueError):
        return "Log"
    if score >= 9.0:
        return "Critical"
    if score >= 7.0:
        return "High"
    if score >= 4.0:
        return "Medium"
    if score > 0.0:
        return "Low"
    return "Log"


def _to_float(value):
    if value is None:
        return None
    try:
        return float(str(value).strip())
    except (TypeError, ValueError):
        return None


def extract_solution(nvt_el, result_el):
    """Решение может быть в <nvt><solution> или в <solution> на уровне result."""
    sol_el = None
    if nvt_el is not None:
        sol_el = nvt_el.find("solution")
    if sol_el is None or not text_or_none(sol_el):
        sol_el = result_el.find("solution")
    return text_or_none(sol_el)


def extract_host(result_el):
    """Возвращает (ip, hostname) из тега <host>."""
    host_el = result_el.find("host")
    if host_el is None:
        return None, None

    ip = clean_text(host_el.text) or None
    hostname = text_or_none(host_el.find("hostname"))
    return ip, hostname


def parse_result(result_el):
    """Извлекает данные одной уязвимости (<result>). None, если нет OID."""
    nvt_el = result_el.find("nvt")
    oid = nvt_el.get("oid") if nvt_el is not None else None
    if not oid:
        return None

    name = None
    if nvt_el is not None:
        name = text_or_none(nvt_el.find("name"))
    if not name:
        name = text_or_none(result_el.find("name"))
    if not name:
        name = "Без названия"

    description = text_or_none(result_el.find("description"))
    solution = extract_solution(nvt_el, result_el)

    cvss_score = _to_float(result_el.findtext("severity"))

    if cvss_score is None and nvt_el is not None:
        severities_el = nvt_el.find("severities")
        if severities_el is not None:
            cvss_score = _to_float(severities_el.get("score"))
            if cvss_score is None:
                # GVM 21+ хранит значение в <severities><severity><score>
                cvss_score = _to_float(severities_el.findtext(".//score"))
        if cvss_score is None:
            cvss_score = _to_float(nvt_el.findtext("cvss_base"))

    if cvss_score is None:
        cvss_score = 0.0
    # Отрицательная severity в OpenVAS означает "лог/ложное срабатывание"
    if cvss_score < 0:
        cvss_score = 0.0

    threat = text_or_none(result_el.find("threat"))
    threat = normalize_threat(threat) if threat else severity_to_threat(cvss_score)

    ip, hostname = extract_host(result_el)
    port = text_or_none(result_el.find("port"))

    return {
        "oid": oid,
        "name": name,
        "description": description,
        "solution": solution,
        "cvss_score": cvss_score,
        "threat_level": threat,
        "ip": ip,
        "hostname": hostname,
        "port": port,
    }


def merge_result(grouped, parsed, stats):
    """Добавляет распарсенный <result> в общую группировку по OID."""
    stats["total_original_results"] += 1

    oid = parsed["oid"]
    entry = grouped.get(oid)
    if entry is None:
        entry = {
            "oid": oid,
            "name": parsed["name"],
            "description": parsed["description"] or "",
            "solution": parsed["solution"] or "",
            "cvss_score": parsed["cvss_score"],
            "threat_level": parsed["threat_level"],
            "_assets": {},  # ip -> {"hostname": ..., "ports": set()}
        }
        grouped[oid] = entry
    else:
        # Берем самое опасное вхождение этого OID: сначала по CVSS,
        # при равенстве — по уровню угрозы.
        current = (entry["cvss_score"], -THREAT_RANK.get(entry["threat_level"], 99))
        candidate = (parsed["cvss_score"], -THREAT_RANK.get(parsed["threat_level"], 99))
        if candidate > current:
            entry["cvss_score"] = parsed["cvss_score"]
            entry["threat_level"] = parsed["threat_level"]
        if not entry["description"] and parsed["description"]:
            entry["description"] = parsed["description"]
        if not entry["solution"] and parsed["solution"]:
            entry["solution"] = parsed["solution"]

    if parsed["ip"]:
        asset = entry["_assets"].setdefault(
            parsed["ip"], {"hostname": parsed["hostname"], "ports": set()}
        )
        if not asset["hostname"] and parsed["hostname"]:
            asset["hostname"] = parsed["hostname"]
        if parsed["port"]:
            asset["ports"].add(parsed["port"])


def _open_xml(filepath):
    """Открывает .xml или .xml.gz в бинарном режиме."""
    if filepath.lower().endswith(".gz"):
        return gzip.open(filepath, "rb")
    return open(filepath, "rb")


def process_file(filepath, grouped, stats):
    """Потоково разбирает один отчет.

    Используется iterparse с очисткой обработанных <result>: отчеты OpenVAS
    на сотни мегабайт не помещаются в память при обычном ET.parse().
    """
    log(f"Обработка файла: {filepath} ...")

    file_count = 0
    seen_result = False
    try:
        with _open_xml(filepath) as fh:
            stack = []
            for event, elem in _safe_iterparse(fh, events=("start", "end")):
                if event == "start":
                    stack.append(elem)
                    continue
                if stack:
                    stack.pop()

                tag = elem.tag
                if tag == "result":
                    seen_result = True
                    parsed = parse_result(elem)
                    if parsed is not None:
                        merge_result(grouped, parsed, stats)
                        file_count += 1
                    # Освобождаем поддерево: дальше оно не нужно
                    elem.clear()
                elif tag in ("creation_time", "timestamp"):
                    # Дата отчета идет до <results>; после первого result
                    # такие теги относятся уже к отдельным находкам.
                    if not seen_result and not stats["scan_date"]:
                        value = clean_text(elem.text)
                        if value:
                            stats["scan_date"] = value
    except ET.ParseError as e:
        warn(f"  [ОШИБКА] Некорректный XML в файле {filepath}: {e}")
        return
    except (OSError, EOFError, gzip.BadGzipFile) as e:
        warn(f"  [ОШИБКА] Не удалось прочитать файл {filepath}: {e}")
        return

    stats["files_processed"] += 1
    log(f"  Обработано записей из файла: {file_count}")


def _ip_sort_key(value):
    """Натуральная сортировка адресов: 10.0.0.2 после 9.0.0.1, а не до."""
    try:
        addr = ipaddress.ip_address(value)
        return (0, addr.version, addr.packed, "")
    except ValueError:
        return (1, 0, b"", value or "")


def _port_sort_key(value):
    """Натуральная сортировка портов: 80/tcp раньше 443/tcp раньше general/tcp."""
    m = re.match(r"^\s*(\d+)", value or "")
    if m:
        return (0, int(m.group(1)), value)
    return (1, 0, value or "")


def build_grouped_data(grouped, stats, source_files, sort_desc=True, min_cvss=None):
    """Собирает финальную структуру данных (аналог JSON-отчета)."""
    vulnerabilities = []
    filtered_out = 0
    for entry in grouped.values():
        if min_cvss is not None and entry["cvss_score"] < min_cvss:
            filtered_out += 1
            continue

        affected_assets = []
        for ip, asset in entry["_assets"].items():
            affected_assets.append(
                {
                    "ip": ip,
                    "hostname": asset["hostname"],
                    "ports": sorted(asset["ports"], key=_port_sort_key),
                }
            )
        affected_assets.sort(key=lambda a: _ip_sort_key(a["ip"]))

        vulnerabilities.append(
            {
                "oid": entry["oid"],
                "name": entry["name"],
                "description": entry["description"],
                "solution": entry["solution"],
                "cvss_score": entry["cvss_score"],
                "threat_level": entry["threat_level"],
                "affected_hosts_count": len(affected_assets),
                "affected_assets": affected_assets,
            }
        )

    if sort_desc:
        # Детерминированный порядок: CVSS ↓, затем уровень угрозы, затем имя
        vulnerabilities.sort(
            key=lambda v: (
                -v["cvss_score"],
                THREAT_RANK.get(v["threat_level"], 99),
                v["name"].lower(),
            )
        )

    if filtered_out:
        log(f"Отфильтровано по --min-cvss: {filtered_out}")

    severity_counts = {level: 0 for level in THREAT_ORDER}
    for v in vulnerabilities:
        severity_counts[normalize_threat(v["threat_level"])] += 1

    return {
        "report_metadata": {
            "scan_date": stats["scan_date"] or datetime.now().isoformat(timespec="seconds"),
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "total_vulnerabilities_grouped": len(vulnerabilities),
            "total_original_results": stats["total_original_results"],
            "severity_counts": severity_counts,
            "min_cvss_filter": min_cvss,
            "source_files": source_files,
        },
        "vulnerabilities": vulnerabilities,
    }


def parse_xml_to_data(input_path, sort_desc=True, recursive=False, min_cvss=None):
    """Полный пайплайн: находит XML, парсит, группирует, возвращает dict-структуру."""
    xml_files = find_xml_files(input_path, recursive=recursive)
    if not xml_files:
        return None

    grouped = {}
    stats = {"total_original_results": 0, "scan_date": None, "files_processed": 0}

    for filepath in xml_files:
        process_file(filepath, grouped, stats)

    if not stats["files_processed"]:
        warn("[ОШИБКА] Ни один из найденных файлов не удалось обработать.")
        return None

    if not grouped:
        warn("[ПРЕДУПРЕЖДЕНИЕ] Не найдено ни одной уязвимости в обработанных файлах.")

    source_files = [os.path.basename(f) for f in xml_files]
    data = build_grouped_data(
        grouped, stats, source_files, sort_desc=sort_desc, min_cvss=min_cvss
    )

    log(f"Обработано файлов: {stats['files_processed']} из {len(xml_files)}")
    log(f"Исходных результатов: {stats['total_original_results']}")
    log(f"Найдено уникальных уязвимостей: {len(grouped)}")

    return data


# ============================================================================
# Часть 2: генерация Word-документа (.docx) из данных
# ============================================================================

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC", size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_column_widths(table, widths_inches):
    """Фиксирует ширины колонок.

    Одних cell.width недостаточно: без w:tblLayout=fixed и корректного
    w:tblGrid Word пересчитывает ширины по содержимому и верстка «плывет».
    """
    table.autofit = False

    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for col, width in zip(grid.findall(qn("w:gridCol")), widths_inches):
            col.set(qn("w:w"), str(int(Inches(width).twips)))

    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)


def repeat_table_header(row):
    """Повторять строку-шапку на каждой странице при переносе таблицы."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_horizontal_rule(paragraph, color="2E75B6", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(clean_text(text) or "-")
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_field(paragraph, instruction):
    """Вставляет поле Word (PAGE, NUMPAGES и т.п.)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_page_numbers(doc):
    """Нижний колонтитул вида «Стр. 3 из 42»."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run_text in ("Стр. ",):
            r = p.add_run(run_text)
            r.font.size = Pt(9)
            r.font.name = "Arial"
        add_field(p, "PAGE")
        r = p.add_run(" из ")
        r.font.size = Pt(9)
        r.font.name = "Arial"
        add_field(p, "NUMPAGES")


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    h1 = styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string("1F1F1F")

    h2 = styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string("1F1F1F")


def add_title_page(doc, meta):
    title = doc.add_heading("Отчет о результатах сканирования уязвимостей", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    rows = [
        ("Дата сканирования: ", str(meta.get("scan_date", "не указана"))),
        ("Отчет сформирован: ", str(meta.get("generated_at", "-"))),
        ("Уникальных уязвимостей: ", str(meta.get("total_vulnerabilities_grouped", 0))),
        ("Всего исходных результатов: ", str(meta.get("total_original_results", 0))),
        ("Источники: ", ", ".join(meta.get("source_files", [])) or "не указаны"),
    ]
    if meta.get("min_cvss_filter") is not None:
        rows.append(("Порог CVSS: ", f"от {meta['min_cvss_filter']:.1f}"))

    for label, value in rows:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(clean_text(value))

    sep = doc.add_paragraph()
    add_horizontal_rule(sep)


def add_summary_table(doc, vulnerabilities):
    counts = {key: 0 for key in THREAT_ORDER}
    for v in vulnerabilities:
        # Неизвестные уровни сводим к "None", иначе они молча выпадали из сводки
        counts[normalize_threat(v.get("threat_level"))] += 1

    doc.add_heading("Сводка по уровням угрозы", level=2)

    table = doc.add_table(rows=2, cols=len(THREAT_ORDER))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [1.15] * len(THREAT_ORDER)
    set_column_widths(table, widths)

    header_cells = table.rows[0].cells
    value_cells = table.rows[1].cells
    for i, level in enumerate(THREAT_ORDER):
        color = THREAT_COLORS[level]
        text_color = "FFFFFF" if level in THREAT_TEXT_LIGHT else "1F1F1F"
        set_cell_text(header_cells[i], RU_THREAT[level], bold=True,
                      color=text_color, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(header_cells[i], color)
        set_cell_borders(header_cells[i])

        set_cell_text(value_cells[i], str(counts[level]), bold=True,
                      size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_borders(value_cells[i])

    doc.add_paragraph()


def add_multiline(doc, label, text):
    """Выводит многострочный текст абзацами, а не одним слипшимся блоком."""
    lines = [ln.strip() for ln in clean_text(text).split("\n")]
    lines = [ln for ln in lines if ln]
    if not lines:
        return

    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(lines[0])
    for line in lines[1:]:
        extra = doc.add_paragraph(line)
        extra.paragraph_format.space_after = Pt(2)


def add_vulnerability_section(doc, vuln, index, total):
    level = normalize_threat(vuln.get("threat_level"))
    color = THREAT_COLORS.get(level, "9E9E9E")

    heading = doc.add_heading(level=2)
    run = heading.add_run(f"{index}. {clean_text(vuln.get('name')) or 'Без названия'}")
    run.font.color.rgb = RGBColor.from_string("1F1F1F")

    # "Плашка" со статусом, CVSS и OID реализована через 1-строчную таблицу с заливкой
    badge_table = doc.add_table(rows=1, cols=3)
    set_column_widths(badge_table, [1.3, 1.6, 2.0])
    cells = badge_table.rows[0].cells
    set_cell_text(cells[0], RU_THREAT.get(level, level), bold=True,
                  color="FFFFFF" if level in THREAT_TEXT_LIGHT else "1F1F1F",
                  size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(cells[0], color)
    set_cell_borders(cells[0])
    set_cell_text(cells[1], f"CVSS: {float(vuln.get('cvss_score') or 0):.1f}",
                  bold=True, size=9)
    set_cell_text(cells[2], f"OID: {vuln.get('oid', '')}", size=8)

    doc.add_paragraph()

    add_multiline(doc, "Описание: ", vuln.get("description"))
    add_multiline(doc, "Решение: ", vuln.get("solution"))

    assets = vuln.get("affected_assets", [])
    if assets:
        doc.add_paragraph().add_run(f"Затронутые хосты ({len(assets)}):").bold = True

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_column_widths(table, [1.6, 2.4, 2.8])

        header = table.rows[0].cells
        for i, label in enumerate(["IP-адрес", "Имя хоста", "Порты"]):
            set_cell_text(header[i], label, bold=True, color="FFFFFF", size=9)
            shade_cell(header[i], "404040")
            set_cell_borders(header[i])
        repeat_table_header(table.rows[0])

        for asset in assets:
            row = table.add_row().cells
            set_cell_text(row[0], asset.get("ip") or "-", size=9)
            set_cell_text(row[1], asset.get("hostname") or "-", size=9)
            set_cell_text(row[2], ", ".join(asset.get("ports") or []) or "-", size=9)
            for cell in row:
                set_cell_borders(cell)
        set_column_widths(table, [1.6, 2.4, 2.8])

    if index < total:
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(12)
        add_horizontal_rule(sep, color="D9D9D9", size=4)


def add_watermark(doc, image_path, width_inches=5.5, opacity=35):
    """
    Добавляет изображение как фоновый водяной знак на каждую страницу документа.

    Реализация: изображение вставляется в заголовок (header) каждой секции
    как якорный объект (wp:anchor) с флагом behindDoc="1" — это помещает его
    за текстом на всех страницах.

    image_path   — путь к файлу изображения (PNG, JPG, BMP и т.д.)
    width_inches — ширина водяного знака в дюймах; высота масштабируется автоматически
    opacity      — непрозрачность 0..100 (35 = слегка видимый фон).
                   Примечание: LibreOffice игнорирует этот параметр при экспорте
                   в PDF; в Microsoft Word он применяется корректно.
    """
    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Файл водяного знака не найден: {image_path}")

    def _t(ns, local):
        return f'{{{ns}}}{local}'

    opacity = max(0, min(100, int(opacity)))
    amt = str(opacity * 1000)  # 0-100 % → 0-100000 (тысячные доли процента)

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        if not header.paragraphs:
            header.add_paragraph()

        # Добавляем картинку через стандартный API — он сам создаёт relationship
        run = header.paragraphs[0].add_run()
        run.add_picture(image_path, width=Inches(width_inches))

        # Находим w:drawing и wp:inline, созданные add_picture
        r_el    = run._r
        drawing = r_el.find(_t(_W,  'drawing'))
        inline  = drawing.find(_t(_WP, 'inline'))

        extent  = inline.find(_t(_WP, 'extent'))
        cx, cy  = extent.get('cx'), extent.get('cy')

        doc_pr  = inline.find(_t(_WP, 'docPr'))
        pic_id  = int(doc_pr.get('id', '1')) if doc_pr is not None else 1

        graphic = inline.find(_t(_A, 'graphic'))

        # Добавляем прозрачность к blip (работает в MS Word)
        blip = graphic.find(
            f'.//{_t(_A, "graphicData")}'
            f'/{_t(_PIC, "pic")}'
            f'/{_t(_PIC, "blipFill")}'
            f'/{_t(_A, "blip")}'
        )
        if blip is not None:
            alpha = etree.SubElement(blip, _t(_A, 'alphaModFix'))
            alpha.set('amt', amt)

        # Абсолютные смещения для центрирования (работает и в Word, и в LibreOffice)
        page_cx = int(section.page_width.emu)  if section.page_width  else 7772400
        page_cy = int(section.page_height.emu) if section.page_height else 10058400
        off_x   = max(0, (page_cx - int(cx)) // 2)
        off_y   = max(0, (page_cy - int(cy)) // 2)

        # Строим wp:anchor — якорный объект за текстом
        anchor = etree.Element(_t(_WP, 'anchor'), nsmap={'wp': _WP, 'a': _A})
        for attr, val in [
            ('distT', '0'), ('distB', '0'), ('distL', '0'), ('distR', '0'),
            ('simplePos',     '0'),
            ('relativeHeight','251658240'),  # z-order позади всего текста
            ('behindDoc',     '1'),          # ← за текстом
            ('locked',        '0'),
            ('layoutInCell',  '1'),
            ('allowOverlap',  '1'),
        ]:
            anchor.set(attr, val)

        sp = etree.SubElement(anchor, _t(_WP, 'simplePos'))
        sp.set('x', '0'); sp.set('y', '0')

        # Горизонтальное центрирование относительно страницы
        pos_h = etree.SubElement(anchor, _t(_WP, 'positionH'))
        pos_h.set('relativeFrom', 'page')
        etree.SubElement(pos_h, _t(_WP, 'posOffset')).text = str(off_x)

        # Вертикальное центрирование относительно страницы
        pos_v = etree.SubElement(anchor, _t(_WP, 'positionV'))
        pos_v.set('relativeFrom', 'page')
        etree.SubElement(pos_v, _t(_WP, 'posOffset')).text = str(off_y)

        ext = etree.SubElement(anchor, _t(_WP, 'extent'))
        ext.set('cx', cx); ext.set('cy', cy)

        eff = etree.SubElement(anchor, _t(_WP, 'effectExtent'))
        for side in ('l', 't', 'r', 'b'):
            eff.set(side, '0')

        etree.SubElement(anchor, _t(_WP, 'wrapNone'))

        dpr = etree.SubElement(anchor, _t(_WP, 'docPr'))
        dpr.set('id',   str(pic_id + 100))
        dpr.set('name', 'Watermark')

        etree.SubElement(anchor, _t(_WP, 'cNvGraphicFramePr'))

        # Перемещаем a:graphic из inline в anchor
        anchor.append(graphic)

        # Заменяем inline на anchor внутри w:drawing
        drawing.remove(inline)
        drawing.append(anchor)


def build_document(data, watermark_path=None, watermark_opacity=35,
                   watermark_width=5.5):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    setup_styles(doc)

    meta = data.get("report_metadata", {})
    vulnerabilities = data.get("vulnerabilities", [])

    add_title_page(doc, meta)
    add_summary_table(doc, vulnerabilities)

    doc.add_heading("Подробное описание уязвимостей", level=1)

    total = len(vulnerabilities)
    if not total:
        doc.add_paragraph("Уязвимости не обнаружены.")

    for i, vuln in enumerate(vulnerabilities, start=1):
        add_vulnerability_section(doc, vuln, i, total)
        if total > 50 and i % 25 == 0:
            log(f"  ... подготовлено разделов: {i} из {total}")

    add_page_numbers(doc)

    if watermark_path:
        log(f"Добавление водяного знака: {watermark_path} ...")
        try:
            add_watermark(doc, watermark_path,
                          width_inches=watermark_width,
                          opacity=watermark_opacity)
        except FileNotFoundError as e:
            warn(f"[ПРЕДУПРЕЖДЕНИЕ] {e} — водяной знак пропущен.")
        except Exception as e:  # битый/неподдерживаемый формат картинки
            warn(f"[ПРЕДУПРЕЖДЕНИЕ] Не удалось добавить водяной знак: {e} — пропущен.")

    return doc


# ============================================================================
# Часть 3: конвертация DOCX -> PDF через LibreOffice
# ============================================================================

def convert_to_pdf(docx_path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        warn("[ПРЕДУПРЕЖДЕНИЕ] LibreOffice (soffice) не найден в PATH. "
             "PDF не создан. Установите LibreOffice (sudo apt install libreoffice) "
             "или откройте .docx и сохраните как PDF вручную.")
        return None

    outdir = os.path.dirname(os.path.abspath(docx_path)) or "."
    pdf_path = os.path.splitext(os.path.abspath(docx_path))[0] + ".pdf"
    before_mtime = os.path.getmtime(pdf_path) if os.path.exists(pdf_path) else None

    # Отдельный профиль: иначе конвертация молча падает, если LibreOffice
    # уже запущен в GUI под тем же пользователем.
    with tempfile.TemporaryDirectory(prefix="soffice-profile-") as profile:
        cmd = [
            soffice,
            f"-env:UserInstallation=file://{profile}",
            "--headless", "--norestore", "--invisible",
            "--convert-to", "pdf", "--outdir", outdir,
            os.path.abspath(docx_path),
        ]
        try:
            proc = subprocess.run(cmd, check=True, capture_output=True, timeout=600)
        except subprocess.CalledProcessError as e:
            warn(f"[ОШИБКА] Конвертация в PDF не удалась: "
                 f"{(e.stderr or b'').decode(errors='ignore').strip()}")
            return None
        except subprocess.TimeoutExpired:
            warn("[ОШИБКА] Конвертация в PDF превысила тайм-аут (600 с).")
            return None

    if not os.path.exists(pdf_path):
        # LibreOffice умеет завершаться с кодом 0, ничего не создав
        warn("[ОШИБКА] LibreOffice завершился без ошибки, но PDF не появился: "
             f"{(proc.stdout or b'').decode(errors='ignore').strip()}")
        return None
    if before_mtime is not None and os.path.getmtime(pdf_path) == before_mtime:
        warn(f"[ОШИБКА] PDF не был перезаписан: {pdf_path}")
        return None

    return pdf_path


# ============================================================================
# main
# ============================================================================

def build_arg_parser():
    parser = argparse.ArgumentParser(
        description=(
            "Группировка уязвимостей из XML-отчетов Greenbone OpenVAS по OID "
            "с выводом в JSON, Word (.docx) и/или PDF."
        )
    )
    parser.add_argument(
        "-i", "--input", required=True,
        help="Путь к XML-файлу (.xml/.xml.gz) или каталогу с отчетами OpenVAS.",
    )
    parser.add_argument(
        "-o", "--output", default="vulnerability_report",
        help="Базовое имя выходных файлов без расширения "
             "(по умолчанию: vulnerability_report).",
    )
    parser.add_argument(
        "-r", "--recursive", action="store_true",
        help="Искать XML-файлы во вложенных каталогах.",
    )
    parser.add_argument(
        "--json", action="store_true",
        help="Сохранить структурированный JSON-отчет (<output>.json).",
    )
    parser.add_argument(
        "--docx", action="store_true",
        help="Сохранить человекочитаемый Word-документ (<output>.docx).",
    )
    parser.add_argument(
        "--pdf", action="store_true",
        help="Сохранить PDF (<output>.pdf). Требует LibreOffice (soffice в PATH).",
    )
    parser.add_argument(
        "--min-cvss", type=float, default=None, metavar="SCORE",
        help="Включать только уязвимости с CVSS не ниже указанного "
             "(например, 4.0 — отсечь информационные записи).",
    )
    parser.add_argument(
        "--no-sort", action="store_true",
        help="Не сортировать уязвимости по CVSS (по умолчанию сортировка по убыванию включена).",
    )
    parser.add_argument(
        "-q", "--quiet", action="store_true",
        help="Не выводить прогресс; ошибки и предупреждения по-прежнему идут в stderr.",
    )
    parser.add_argument(
        "--watermark", metavar="IMAGE",
        help="Путь к файлу изображения (PNG/JPG/BMP), который будет добавлен "
             "как фоновый водяной знак на каждую страницу DOCX/PDF.",
    )
    parser.add_argument(
        "--watermark-opacity", type=int, default=35, metavar="0-100",
        help="Непрозрачность водяного знака: 0 = невидимый, 100 = сплошной "
             "(по умолчанию: 35). Работает в MS Word; LibreOffice игнорирует.",
    )
    parser.add_argument(
        "--watermark-width", type=float, default=5.5, metavar="INCHES",
        help="Ширина водяного знака в дюймах (по умолчанию: 5.5). "
             "Высота масштабируется автоматически.",
    )
    return parser


def validate_args(parser, args):
    if not 0 <= args.watermark_opacity <= 100:
        parser.error("--watermark-opacity должен быть в диапазоне 0..100")
    if args.watermark_width <= 0:
        parser.error("--watermark-width должен быть больше нуля")
    if args.min_cvss is not None and not 0 <= args.min_cvss <= 10:
        parser.error("--min-cvss должен быть в диапазоне 0..10")


def ensure_output_dir(path):
    """Создает каталог для выходного файла, если он задан через -o dir/name."""
    directory = os.path.dirname(os.path.abspath(path))
    if directory and not os.path.isdir(directory):
        os.makedirs(directory, exist_ok=True)


def write_json(data, json_path):
    try:
        ensure_output_dir(json_path)
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return json_path
    except OSError as e:
        warn(f"[ОШИБКА] Не удалось записать {json_path}: {e}")
        return None


def run(args):
    global QUIET
    QUIET = args.quiet

    if not HAVE_DEFUSEDXML:
        log("[ИНФО] Пакет defusedxml не установлен — разбор XML без защиты "
            "от entity-бомб (pip install defusedxml).")

    data = parse_xml_to_data(
        args.input,
        sort_desc=not args.no_sort,
        recursive=args.recursive,
        min_cvss=args.min_cvss,
    )
    if data is None:
        warn("Не найдено файлов для обработки. Завершение.")
        return 1

    produced = []
    failed = []

    if args.json:
        path = write_json(data, f"{args.output}.json")
        if path:
            produced.append(path)
        else:
            failed.append("JSON")

    docx_path = f"{args.output}.docx"
    docx_created_for_pdf_only = False

    if args.docx or args.pdf:
        log("Формирование Word-документа ...")
        doc = build_document(
            data,
            watermark_path=args.watermark,
            watermark_opacity=args.watermark_opacity,
            watermark_width=args.watermark_width,
        )
        try:
            ensure_output_dir(docx_path)
            doc.save(docx_path)
        except OSError as e:
            warn(f"[ОШИБКА] Не удалось сохранить {docx_path}: {e}")
            docx_path = None
            if args.docx:
                failed.append("DOCX")
            if args.pdf:
                failed.append("PDF")

        if docx_path:
            if args.docx:
                produced.append(docx_path)
            else:
                docx_created_for_pdf_only = True

    if args.pdf and docx_path:
        log("Конвертация в PDF ...")
        pdf_path = convert_to_pdf(docx_path)
        if pdf_path:
            produced.append(pdf_path)
        else:
            failed.append("PDF")

        # Промежуточный .docx удаляем только при успешной конвертации,
        # иначе пользователь остался бы вообще без результата.
        if docx_created_for_pdf_only:
            if pdf_path:
                try:
                    os.remove(docx_path)
                except OSError as e:
                    warn(f"[ПРЕДУПРЕЖДЕНИЕ] Не удалось удалить временный "
                         f"{docx_path}: {e}")
            else:
                warn(f"[ПРЕДУПРЕЖДЕНИЕ] PDF не создан, промежуточный файл "
                     f"сохранен: {docx_path}")
                produced.append(docx_path)

    log("")
    log("=== Готово ===")
    for path in produced:
        log(f"  -> {path}")

    if not produced:
        warn("Ничего не было сохранено (проверьте сообщения об ошибках выше).")
        return 1
    if failed:
        # Часть запрошенных форматов не получена — сообщаем это кодом возврата,
        # чтобы вызывающий скрипт не считал запуск успешным.
        warn(f"[ОШИБКА] Не удалось создать: {', '.join(failed)}")
        return 1
    return 0


def main():
    parser = build_arg_parser()
    args = parser.parse_args()
    validate_args(parser, args)

    # Если ни один формат не указан явно — по умолчанию делаем Word-документ,
    # как наиболее человекочитаемый вариант.
    if not (args.json or args.docx or args.pdf):
        if not args.quiet:
            print("Формат вывода не указан (--json/--docx/--pdf), "
                  "по умолчанию будет создан Word-документ (--docx).", flush=True)
        args.docx = True

    try:
        return run(args)
    except KeyboardInterrupt:
        warn("\nПрервано пользователем.")
        return 130


if __name__ == "__main__":
    sys.exit(main())
